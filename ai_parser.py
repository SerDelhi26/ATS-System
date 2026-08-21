import os
import re
import io
import json
import base64
import requests
from dotenv import load_dotenv
import streamlit as st

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

load_dotenv()

def get_gemini_keys() -> list[str]:
    """Retrieves all configured Gemini API keys (supports individual user keys & pool)."""
    keys = []
    # Check individual user variables first
    for i in range(1, 10):
        k = os.getenv(f"GEMINI_API_KEY_USER{i}")
        if not k:
            try:
                k = st.secrets.get(f"GEMINI_API_KEY_USER{i}")
            except Exception:
                pass
        if k and str(k).strip() and str(k).strip() not in keys:
            keys.append(str(k).strip())
            
    # Check pool string
    env_keys = os.getenv("GEMINI_API_KEYS") or os.getenv("GEMINI_API_KEY") or ""
    if not env_keys:
        try:
            env_keys = st.secrets.get("GEMINI_API_KEYS", "") or st.secrets.get("GEMINI_API_KEY", "")
        except Exception:
            pass

    for k in str(env_keys).split(","):
        k_clean = k.strip()
        if k_clean and k_clean not in keys:
            keys.append(k_clean)
    return keys


def get_groq_keys() -> list[str]:
    """Retrieves all configured Groq API keys (supports individual user keys & pool)."""
    keys = []
    for i in range(1, 10):
        k = os.getenv(f"GROQ_API_KEY_USER{i}")
        if not k:
            try:
                k = st.secrets.get(f"GROQ_API_KEY_USER{i}")
            except Exception:
                pass
        if k and str(k).strip() and str(k).strip() not in keys:
            keys.append(str(k).strip())

    env_keys = os.getenv("GROQ_API_KEYS") or os.getenv("GROQ_API_KEY") or ""
    if not env_keys:
        try:
            env_keys = st.secrets.get("GROQ_API_KEYS", "") or st.secrets.get("GROQ_API_KEY", "")
        except Exception:
            pass

    for k in str(env_keys).split(","):
        k_clean = k.strip()
        if k_clean and k_clean not in keys:
            keys.append(k_clean)
    return keys


def get_openrouter_keys() -> list[str]:
    """Retrieves all configured OpenRouter API keys (supports individual user keys & pool)."""
    keys = []
    for i in range(1, 10):
        k = os.getenv(f"OPENROUTER_API_KEY_USER{i}")
        if not k:
            try:
                k = st.secrets.get(f"OPENROUTER_API_KEY_USER{i}")
            except Exception:
                pass
        if k and str(k).strip() and str(k).strip() not in keys:
            keys.append(str(k).strip())

    env_keys = os.getenv("OPENROUTER_API_KEYS") or os.getenv("OPENROUTER_API_KEY") or ""
    if not env_keys:
        try:
            env_keys = st.secrets.get("OPENROUTER_API_KEYS", "") or st.secrets.get("OPENROUTER_API_KEY", "")
        except Exception:
            pass

    for k in str(env_keys).split(","):
        k_clean = k.strip()
        if k_clean and k_clean not in keys:
            keys.append(k_clean)
    return keys


def clean_phone(phone_str: str) -> str:
    """Extracts the last 10 digits from a phone string."""
    if not phone_str:
        return ""
    digits = re.sub(r'\D', '', str(phone_str))
    return digits[-10:] if len(digits) >= 10 else digits


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text content directly from PDF bytes in milliseconds."""
    try:
        if pypdf is None:
            return ""
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        extracted_pages = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                extracted_pages.append(t)
        return "\n".join(extracted_pages).strip()
    except Exception:
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text content directly from Word .docx bytes."""
    try:
        if docx is None:
            return ""
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception:
        return ""


def sanitize_parsed_output(parsed: dict) -> dict:
    """Sanitizes and normalizes extracted fields. Notice period, notice negotiable, and remarks are excluded."""
    return {
        "first_name": str(parsed.get("first_name", "")).strip(),
        "last_name": str(parsed.get("last_name", "")).strip(),
        "email": str(parsed.get("email", "")).strip().lower(),
        "mobile_no": clean_phone(parsed.get("mobile_no", "")),
        "alternate_mobile": clean_phone(parsed.get("alternate_mobile", "")),
        "current_location": str(parsed.get("current_location", "")).strip(),
        "experience_years": max(0, min(40, int(parsed.get("experience_years", 0) or 0))),
        "experience_months": max(0, min(11, int(parsed.get("experience_months", 0) or 0))),
        "qualification": str(parsed.get("qualification", "")).strip(),
        "education_details": str(parsed.get("education_details", "")).strip(),
        "current_company": str(parsed.get("current_company", "")).strip(),
        "current_designation": str(parsed.get("current_designation", "")).strip(),
        "current_ctc": float(parsed.get("current_ctc", 0.0) or 0.0),
        "expected_ctc": float(parsed.get("expected_ctc", 0.0) or 0.0),
        "skills": str(parsed.get("skills", "")).strip(),
    }


def _call_gemini_api(api_key: str, model: str, payload: dict) -> tuple[bool, dict, str]:
    """Makes a single call to the Google Gemini generateContent API."""
    endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    try:
        response = requests.post(endpoint, headers=headers, json=payload, timeout=40)
        if response.status_code == 429:
            return False, {}, "RATE_LIMIT_429"
        if response.status_code != 200:
            return False, {}, f"Gemini Error ({response.status_code}): {response.text[:200]}"

        res_json = response.json()
        candidates = res_json.get("candidates", [])
        if not candidates:
            return False, {}, "No candidates returned"

        raw_text = candidates[0]["content"]["parts"][0]["text"].strip()
        if raw_text.startswith("```json"):
            raw_text = raw_text[7:]
        if raw_text.startswith("```"):
            raw_text = raw_text[3:]
        if raw_text.endswith("```"):
            raw_text = raw_text[:-3]

        parsed = json.loads(raw_text.strip())
        return True, sanitize_parsed_output(parsed), "Success"
    except requests.exceptions.Timeout:
        return False, {}, "TIMEOUT"
    except Exception as e:
        return False, {}, str(e)


def _call_groq_api(api_key: str, model: str, system_prompt: str, resume_text: str) -> tuple[bool, dict, str]:
    """Makes a call to Groq Cloud OpenAI-compatible chat API."""
    endpoint = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"RESUME TEXT:\n{resume_text}"}
        ],
        "response_format": {"type": "json_object"},
        "temperature": 0.1
    }
    try:
        response = requests.post(endpoint, headers=headers, json=payload, timeout=30)
        if response.status_code == 429:
            return False, {}, "RATE_LIMIT_429"
        if response.status_code != 200:
            return False, {}, f"Groq Error ({response.status_code}): {response.text[:200]}"

        res_json = response.json()
        content = res_json["choices"][0]["message"]["content"]
        parsed = json.loads(content.strip())
        return True, sanitize_parsed_output(parsed), "Success"
    except requests.exceptions.Timeout:
        return False, {}, "TIMEOUT"
    except Exception as e:
        return False, {}, str(e)


def _call_openrouter_api(api_key: str, model: str, system_prompt: str, resume_text: str) -> tuple[bool, dict, str]:
    """Makes a call to OpenRouter OpenAI-compatible chat API."""
    endpoint = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://localhost:8501",
        "X-Title": "ATS Resume Parser"
    }
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"RESUME TEXT:\n{resume_text}"}
        ],
        "response_format": {"type": "json_object"},
        "temperature": 0.1
    }
    try:
        response = requests.post(endpoint, headers=headers, json=payload, timeout=35)
        if response.status_code == 429:
            return False, {}, "RATE_LIMIT_429"
        if response.status_code != 200:
            return False, {}, f"OpenRouter Error ({response.status_code}): {response.text[:200]}"

        res_json = response.json()
        content = res_json["choices"][0]["message"]["content"]
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]

        parsed = json.loads(content.strip())
        return True, sanitize_parsed_output(parsed), "Success"
    except requests.exceptions.Timeout:
        return False, {}, "TIMEOUT"
    except Exception as e:
        return False, {}, str(e)


def parse_resume_with_ai(file_bytes: bytes, filename: str, mime_type: str = None) -> tuple[bool, dict, str]:
    """
    Parses a candidate resume with Multi-Key & Multi-Provider load balancing and failover:
    Gemini Key Pool -> Groq Key Pool -> OpenRouter Key Pool.
    """
    if not file_bytes:
        return False, {}, "No file content provided."

    gemini_keys = get_gemini_keys()
    groq_keys = get_groq_keys()
    openrouter_keys = get_openrouter_keys()

    if not gemini_keys and not groq_keys and not openrouter_keys:
        return False, {}, "No AI API keys configured. Please add GEMINI_API_KEY, GROQ_API_KEY, or OPENROUTER_API_KEY in .env or secrets.toml."

    system_prompt = """You are an expert AI Resume Parser for an enterprise ATS.
Extract candidate information from the resume into this exact JSON structure:
{
  "first_name": "Candidate's First name",
  "last_name": "Candidate's Last name (or empty string if none)",
  "email": "Email address or empty string",
  "mobile_no": "10-digit primary mobile number or empty string",
  "alternate_mobile": "Secondary contact number or empty string",
  "current_location": "Current city / state or location",
  "experience_years": 0,
  "experience_months": 0,
  "qualification": "Highest degree/qualification (e.g. B.Tech, MCA, MBA, B.Sc, Diploma, M.Tech, Graduate)",
  "education_details": "Summary of education, degrees, colleges and years",
  "current_company": "Current or most recent employer/company name",
  "current_designation": "Current or most recent job title/designation",
  "current_ctc": 0.0,
  "expected_ctc": 0.0,
  "skills": "Comma-separated key skills, frameworks, and technologies"
}

Important Rules:
1. Extract first and last names cleanly (remove titles like Mr., Ms., Dr.).
2. Calculate total professional experience accurately in full years (integer) and remaining months (0-11 integer).
3. If CTC is not explicitly stated, return 0.0.
4. Output valid JSON only.
"""

    ext = os.path.splitext(filename or "")[1].lower()

    # 1. Fast Local Text Extraction
    extracted_text = ""
    if ext == ".pdf" or (mime_type and "pdf" in mime_type.lower()):
        extracted_text = extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        extracted_text = extract_text_from_docx(file_bytes)
    else:
        try:
            extracted_text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            extracted_text = ""

    # Build Gemini Payload
    gemini_parts = [{"text": system_prompt}]
    if len(extracted_text) >= 20:
        gemini_parts.append({"text": f"DOCUMENT FILENAME: {filename}\n\nRESUME CONTENT:\n{extracted_text[:20000]}"})
    elif ext == ".pdf" or (mime_type and "pdf" in mime_type.lower()):
        b64_data = base64.b64encode(file_bytes).decode("utf-8")
        gemini_parts.append({
            "inlineData": {
                "mimeType": "application/pdf",
                "data": b64_data
            }
        })
    else:
        gemini_parts.append({"text": f"DOCUMENT: {filename}\n{file_bytes.decode('latin-1', errors='ignore')[:15000]}"})

    gemini_payload = {
        "contents": [{"parts": gemini_parts}],
        "generationConfig": {
            "responseMimeType": "application/json",
            "temperature": 0.1
        }
    }

    errors = []

    # -------------------------------------------------------------
    # 1. Try Gemini Key Pool & Models
    # -------------------------------------------------------------
    gemini_models = ["gemini-3.6-flash", "gemini-3.5-flash-lite", "gemini-3.5-flash"]
    for idx, key in enumerate(gemini_keys, 1):
        for model in gemini_models:
            success, data, msg = _call_gemini_api(key, model, gemini_payload)
            if success:
                return True, data, "Resume parsed successfully via Gemini AI!"
            if msg == "RATE_LIMIT_429":
                errors.append(f"Gemini Key #{idx} rate-limited (429).")
                break # Switch to next key in pool immediately
            else:
                errors.append(f"Gemini ({model}): {msg}")

    # -------------------------------------------------------------
    # 2. Try Groq Key Pool (if Gemini keys exhausted or failed)
    # -------------------------------------------------------------
    if groq_keys and extracted_text:
        groq_models = ["llama-3.3-70b-versatile", "llama-3.1-8b-instant"]
        for idx, key in enumerate(groq_keys, 1):
            for model in groq_models:
                success, data, msg = _call_groq_api(key, model, system_prompt, extracted_text)
                if success:
                    return True, data, "Resume parsed successfully via Groq AI!"
                if msg == "RATE_LIMIT_429":
                    errors.append(f"Groq Key #{idx} rate-limited (429).")
                    break # Switch to next Groq key
                else:
                    errors.append(f"Groq ({model}): {msg}")

    # -------------------------------------------------------------
    # 3. Try OpenRouter Key Pool (if Gemini & Groq exhausted or failed)
    # -------------------------------------------------------------
    if openrouter_keys and extracted_text:
        openrouter_models = [
            "meta-llama/llama-3.3-70b-instruct:free",
            "google/gemini-2.0-flash-exp:free",
            "mistralai/mistral-7b-instruct:free"
        ]
        for idx, key in enumerate(openrouter_keys, 1):
            for model in openrouter_models:
                success, data, msg = _call_openrouter_api(key, model, system_prompt, extracted_text)
                if success:
                    return True, data, "Resume parsed successfully via OpenRouter AI!"
                if msg == "RATE_LIMIT_429":
                    errors.append(f"OpenRouter Key #{idx} rate-limited (429).")
                    break
                else:
                    errors.append(f"OpenRouter ({model}): {msg}")

    err_summary = " | ".join(errors[-3:]) if errors else "All AI providers failed."
    return False, {}, f"AI Parsing Failed: {err_summary}"
